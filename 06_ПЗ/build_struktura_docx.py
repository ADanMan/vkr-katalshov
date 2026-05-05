#!/usr/bin/env python3
"""
build_struktura_docx.py — сборка документа «Структура и содержание ВКР»
по ГОСТ 7.32-2017 (без К4-шаблона; чистый формат).

Использование:
    python3 build_struktura_docx.py \
        --src draft/02_Struktura_VKR.md \
        --out Struktura_VKR.docx

Стиль:
* Times New Roman 14 pt, межстрочный 1,5;
* поля: левое 30 мм, правое 10 мм, верхнее/нижнее 20 мм;
* абзацный отступ 1,25 см, выравнивание по ширине;
* заголовки h1: прописными, по центру, жирный;
* заголовки h2: с прописной по абзацу, жирный;
* подписи таблиц: «Таблица N — Название» сверху, без отступа.

В документе формируется простой титульный «лист»: шапка министерства,
факультет/кафедра, название документа, тема ВКР, студент, руководитель,
город и год — по требованиям ГОСТ 7.32-2017 без К4-шаблона.

Затем — содержательная часть из Markdown-черновика с поддержкой
inline-форматирования **bold**, *italic*, `code`.
"""

from __future__ import annotations

import argparse
import logging
import re
import sys
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

FONT = "Times New Roman"
FONT_SIZE = 14
LINE_SPACING = 1.5
INDENT_CM = 1.25


# ──────────────────────────────────────────────────────────────────────
# Низкоуровневые утилиты docx
# ──────────────────────────────────────────────────────────────────────


def _set_run_font(run: Any, *, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = FONT
    run.font.size = Pt(FONT_SIZE)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), FONT)


def _set_para(p: Any, *, indent: bool, alignment: str) -> None:
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if indent:
        pf.first_line_indent = Cm(INDENT_CM)
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map[alignment]


_INLINE = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def _add_inline_runs(p: Any, text: str) -> None:
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            _set_run_font(r, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1])
            _set_run_font(r, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            _set_run_font(r)
            r.font.name = "Consolas"
        else:
            r = p.add_run(part)
            _set_run_font(r)


def _set_cell_borders(cell: Any) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        b = tcBorders.find(qn(f"w:{edge}"))
        if b is None:
            b = OxmlElement(f"w:{edge}")
            tcBorders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        b.set(qn("w:space"), "0")


# ──────────────────────────────────────────────────────────────────────
# Настройки документа: поля и стиль Normal
# ──────────────────────────────────────────────────────────────────────


def _setup_doc(doc: Document) -> None:
    # Поля по ГОСТ 7.32-2017
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
    # Стиль Normal: TNR 14, 1,5
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(FONT_SIZE)
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


# ──────────────────────────────────────────────────────────────────────
# Титульный лист в стиле ГОСТ 7.32 (без К4-бланка)
# ──────────────────────────────────────────────────────────────────────


def _add_centered(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    _set_para(p, indent=False, alignment="center")
    r = p.add_run(text)
    _set_run_font(r, bold=bold, italic=italic)


def _add_blank(doc: Document, n: int = 1) -> None:
    for _ in range(n):
        doc.add_paragraph()


def add_title_page(doc: Document, meta: dict[str, Any]) -> None:
    # Шапка
    for line in (
        "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ",
        "РОССИЙСКОЙ ФЕДЕРАЦИИ",
    ):
        _add_centered(doc, line, bold=True)
    _add_centered(
        doc,
        "Федеральное государственное бюджетное образовательное учреждение высшего образования",
    )
    _add_centered(
        doc,
        "«Московский государственный технический университет имени Н.Э. Баумана (национальный исследовательский университет)»",
        bold=True,
    )
    _add_centered(doc, "(МГТУ им. Н.Э. Баумана)", bold=True)
    _add_blank(doc)
    _add_centered(doc, meta.get("branch", "Мытищинский филиал"))
    _add_centered(doc, f"Факультет {meta.get('faculty', '')}")
    _add_centered(doc, f"Кафедра {meta.get('department', '')}")

    _add_blank(doc, 4)

    _add_centered(doc, "СТРУКТУРА И СОДЕРЖАНИЕ", bold=True)
    _add_centered(doc, "выпускной квалификационной работы бакалавра", italic=True)

    _add_blank(doc, 2)
    _add_centered(doc, "на тему:")
    _add_centered(
        doc,
        f"«{meta.get('topic', '').strip()}»",
        bold=True,
    )

    _add_blank(doc, 4)

    # Студент / руководитель — выровнены по абзацу
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    rows = [
        ("Направление:", meta.get("direction", "")),
        ("Профиль:", meta.get("profile", "")),
        ("Студент:", f"{meta.get('student', '')}, группа {meta.get('group', '')}"),
        (
            "Руководитель ВКР:",
            f"{meta.get('supervisor_title', '')} {_short_name(meta.get('supervisor_full', ''))}",
        ),
    ]
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.paragraphs[0].text = ""
        c1.paragraphs[0].text = ""
        r0 = c0.paragraphs[0].add_run(label)
        r1 = c1.paragraphs[0].add_run(value)
        _set_run_font(r0, bold=True)
        _set_run_font(r1)
        c0.width = Cm(5)
        c1.width = Cm(11)

    _add_blank(doc, 4)

    _add_centered(doc, f"г. {meta.get('city', 'Мытищи')}, {meta.get('year', 2026)} г.")

    doc.add_page_break()


def _short_name(full: str) -> str:
    parts = full.split()
    if len(parts) >= 3:
        return f"{parts[1][0]}.{parts[2][0]}. {parts[0]}"
    if len(parts) == 2:
        return f"{parts[1][0]}. {parts[0]}"
    return full


# ──────────────────────────────────────────────────────────────────────
# Парсер Markdown
# ──────────────────────────────────────────────────────────────────────


def parse_markdown(src: Path) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    text = src.read_text(encoding="utf-8")
    fm: dict[str, Any] = {}
    body = text
    if text.startswith("---\n"):
        end = text.find("\n---", 4)
        if end > 0:
            fm = yaml.safe_load(text[4:end]) or {}
            body = text[end + len("\n---") :]

    blocks: list[dict[str, Any]] = []
    lines = body.splitlines()
    i = 0
    pending_caption: str | None = None
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            blocks.append({"kind": "h1", "text": line[2:].strip()})
            i += 1
            continue
        if line.startswith("## "):
            blocks.append({"kind": "h2", "text": line[3:].strip()})
            i += 1
            continue
        m = re.match(r"^\*\*(Таблица\s+\d+\s*[—–-]\s*.+?)\*\*\s*$", line)
        if m:
            pending_caption = m.group(1).strip()
            i += 1
            continue
        if line.lstrip().startswith("|"):
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 1
            if i < len(lines) and re.match(r"^\s*\|[-: |]+\|\s*$", lines[i]):
                i += 1
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            blocks.append({
                "kind": "table",
                "caption": pending_caption,
                "headers": headers,
                "rows": rows,
            })
            pending_caption = None
            continue
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith("#")
            or lines[i].lstrip().startswith("|")
            or lines[i].lstrip().startswith("**Таблица")
        ):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"kind": "p", "text": " ".join(s.strip() for s in para_lines)})
    return fm, blocks


# ──────────────────────────────────────────────────────────────────────
# Контентные блоки
# ──────────────────────────────────────────────────────────────────────


def _add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para(p, indent=False, alignment="center")
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.keep_with_next = True
    r = p.add_run(text.upper())
    _set_run_font(r, bold=True)


def _add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para(p, indent=True, alignment="left")
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.keep_with_next = True
    r = p.add_run(text)
    _set_run_font(r, bold=True)


def _add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _set_para(p, indent=True, alignment="justify")
    _add_inline_runs(p, text)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str | None = None) -> None:
    if caption:
        cap_p = doc.add_paragraph()
        _set_para(cap_p, indent=False, alignment="left")
        _add_inline_runs(cap_p, caption)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        for row in table.rows:
            for cell in row.cells:
                _set_cell_borders(cell)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].add_run(h)
        _set_run_font(r, bold=True)

    for ri, row in enumerate(rows, start=1):
        for ci, value in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(cell.paragraphs[0], value)


def append_content(doc: Document, blocks: list[dict[str, Any]]) -> None:
    for blk in blocks:
        kind = blk["kind"]
        if kind == "h1":
            _add_h1(doc, blk["text"])
        elif kind == "h2":
            _add_h2(doc, blk["text"])
        elif kind == "p":
            _add_paragraph(doc, blk["text"])
        elif kind == "table":
            _add_table(doc, blk["headers"], blk["rows"], blk.get("caption"))


# ──────────────────────────────────────────────────────────────────────
# main
# ──────────────────────────────────────────────────────────────────────


def build(src: Path, out: Path) -> None:
    fm, blocks = parse_markdown(src)
    doc = Document()
    _setup_doc(doc)
    add_title_page(doc, fm)
    append_content(doc, blocks)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("Сохранён %s (%d байт)", out, out.stat().st_size)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--src", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args(argv)
    if not args.src.exists():
        print(f"Source not found: {args.src}", file=sys.stderr)
        return 2
    build(args.src, args.out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
