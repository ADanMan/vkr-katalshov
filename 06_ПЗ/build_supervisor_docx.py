#!/usr/bin/env python3
"""
build_supervisor_docx.py — собирает документ для научрука поверх
исходного шаблона К4 МФ МГТУ.

Подход: НЕ перерисовываем титульный лист, а аккуратно заполняем
плейсхолдеры "_____" в существующих параграфах шаблона через
прямое присваивание `run.text = ...`. Это сохраняет 100 % исходной
вёрстки: шрифт 22 pt для «РАСЧЕТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА», 20 pt
курсив жирный для темы, табуляции и подписи под линиями подписей.

Отличия от предыдущей версии:
* Не подменяются заголовки «РАСЧЕТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА К ВКР НА
  ТЕМУ:» — они остаются как в К4.
* Тема пишется в пять отведённых строк (параграфы 13-17 шаблона).
* Студент: добавляется группа и фамилия, прочерки сохраняются.
* Руководитель ВКР: добавляется звание и фамилия.
* Консультанты и нормоконтролёр оставляются с прочерками
  (заполнятся вручную при сдаче).
* Год — заменяется «20      г.» → «2026 г.».
* Из шаблона остаётся только первая страница; страницы Задания,
  Календарного плана и Направления удаляются.

Использование:
    python3 build_supervisor_docx.py \
        --src draft/01_Otchet_dlya_nauchruka.md \
        --template ../01_Источники/03_МГТУ/10_blank_vkrb_K4.doc \
        --out Otchet_dlya_nauchruka.docx
"""

from __future__ import annotations

import argparse
import logging
import re
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# Стиль ОСНОВНОГО ТЕКСТА (после титула)
BODY_FONT = "Times New Roman"
BODY_FONT_SIZE = 14
BODY_LINE_SPACING = 1.5
BODY_INDENT_CM = 1.25


# ----------------------------------------------------------------------
# Конвертация .doc → .docx через soffice (если шаблон в .doc)
# ----------------------------------------------------------------------


def _ensure_docx(template: Path) -> Path:
    if template.suffix.lower() == ".docx":
        return template
    if template.suffix.lower() != ".doc":
        raise SystemExit(f"Поддерживаются только .doc / .docx, получен {template.suffix}")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice is None:
        raise SystemExit("Нужен LibreOffice (soffice) для конвертации .doc → .docx.")
    out_dir = Path(tempfile.mkdtemp(prefix="k4_"))
    logger.info("Конвертация %s → %s", template.name, out_dir)
    subprocess.run(
        [soffice, "--headless", "--convert-to", "docx", "--outdir", str(out_dir), str(template)],
        check=True,
        capture_output=True,
    )
    out = out_dir / (template.stem + ".docx")
    if not out.exists():
        raise SystemExit(f"Не удалось сконвертировать {template}")
    return out


# ----------------------------------------------------------------------
# Удаление лишних страниц / таблиц
# ----------------------------------------------------------------------


def _delete_paragraph(paragraph: Any) -> None:
    p = paragraph._element
    p.getparent().remove(p)


def _delete_table(table: Any) -> None:
    t = table._element
    t.getparent().remove(t)


def _strip_to_title_page(doc: Document, title_end_paragraph: int = 39) -> None:
    """Оставить только параграфы 0..title_end_paragraph; удалить все
    последующие параграфы и все таблицы."""
    paragraphs = list(doc.paragraphs)
    for p in paragraphs[title_end_paragraph + 1 :]:
        try:
            _delete_paragraph(p)
        except Exception:
            pass
    for t in list(doc.tables):
        try:
            _delete_table(t)
        except Exception:
            pass


# ----------------------------------------------------------------------
# Заполнение плейсхолдеров К4
# ----------------------------------------------------------------------


def _wrap_topic(topic: str, *, max_lines: int = 5, max_chars: int = 60) -> list[str]:
    words = topic.strip().split()
    lines: list[str] = []
    cur = ""
    for w in words:
        candidate = (cur + " " + w).strip() if cur else w
        if len(candidate) > max_chars and cur:
            lines.append(cur)
            cur = w
        else:
            cur = candidate
    if cur:
        lines.append(cur)
    if len(lines) > max_lines:
        # Сольём хвост в последнюю строку
        head = lines[: max_lines - 1]
        tail = " ".join(lines[max_lines - 1 :])
        lines = head + [tail]
    return lines


def _short_name(full: str) -> str:
    """«Катальшов Данила Алексеевич» → «Катальшов Д.А.»."""
    parts = full.split()
    if len(parts) >= 3:
        return f"{parts[0]} {parts[1][0]}.{parts[2][0]}."
    if len(parts) == 2:
        return f"{parts[0]} {parts[1][0]}."
    return full


def _initials_first(full: str) -> str:
    """«Комаров Евгений Геннадьевич» → «Е.Г. Комаров»."""
    parts = full.split()
    if len(parts) >= 3:
        return f"{parts[1][0]}.{parts[2][0]}. {parts[0]}"
    if len(parts) == 2:
        return f"{parts[1][0]}. {parts[0]}"
    return full


def fill_title_page_k4(doc: Document, meta: dict[str, Any]) -> None:
    """Заполнить плейсхолдеры в первой странице шаблона К4.

    Главный принцип — ничего не пересоздаём, лишь меняем `run.text`
    в существующих параграфах. Все стили (шрифт, размер, bold/italic,
    табуляция, alignment, sectPr) сохраняются.
    """
    paras = doc.paragraphs

    # ─── № 2 — ФАКУЛЬТЕТ ─────────────────────────────────────────────
    # Один run: "ФАКУЛЬТЕТ _____________..."
    if len(paras[2].runs) >= 1:
        paras[2].runs[0].text = f"ФАКУЛЬТЕТ {meta.get('faculty', 'Космический')}"

    # ─── № 4 — КАФЕДРА ───────────────────────────────────────────────
    # Два runа: "КАФЕДРА _" и "________________________________"
    department = meta.get(
        "department", "К2 «Информационно-измерительная техника и технологии»"
    )
    if len(paras[4].runs) >= 1:
        paras[4].runs[0].text = f"КАФЕДРА {department}"
    if len(paras[4].runs) >= 2:
        paras[4].runs[1].text = ""

    # ─── № 8, 10, 12 — заголовки. НЕ ТРОГАЕМ. ────────────────────────
    # «РАСЧЕТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА», «К ВЫПУСКНОЙ...», «НА ТЕМУ:»

    # ─── № 13-17 — пять строк темы ───────────────────────────────────
    topic = meta.get("topic", "").strip()
    topic_lines = _wrap_topic(topic, max_lines=5, max_chars=58)
    for idx, line in enumerate(topic_lines):
        target = 13 + idx
        if target > 17:
            break
        if paras[target].runs:
            paras[target].runs[0].text = line
    # Оставшиеся пустые строки — очищаем (вместо прочерков)
    for idx in range(len(topic_lines), 5):
        target = 13 + idx
        if paras[target].runs:
            paras[target].runs[0].text = ""

    # ─── № 21 — Студент ──────────────────────────────────────────────
    # «Студент ________________\t\t\t\t» + «_________________  ____________________»
    student_full = meta.get("student", "Катальшов Данила Алексеевич")
    student_short = _short_name(student_full)
    group = meta.get("group", "К2-81Б")
    if len(paras[21].runs) >= 1:
        paras[21].runs[0].text = f"Студент {group}\t\t\t\t"
    if len(paras[21].runs) >= 2:
        paras[21].runs[1].text = f"_________________  {student_short} "

    # ─── № 24 — Руководитель ВКР ─────────────────────────────────────
    supervisor_full = meta.get("supervisor_full", "Комаров Евгений Геннадьевич")
    supervisor_title = meta.get("supervisor_title", "к.т.н., доцент")
    supervisor_short = _initials_first(supervisor_full)
    if len(paras[24].runs) >= 1:
        paras[24].runs[0].text = f"Руководитель ВКР, {supervisor_title}\t\t\t"
    if len(paras[24].runs) >= 2:
        paras[24].runs[1].text = f"_________________  {supervisor_short} "

    # ─── № 27, 30 — Консультанты — оставляем прочерки в шаблоне ──────
    # ─── № 33 — Нормоконтролёр — оставляем прочерки ──────────────────

    # ─── № 39 — год ──────────────────────────────────────────────────
    # «20      г.» (italic) → «г. Мытищи, 2026 г.»
    city = meta.get("city", "Мытищи")
    year = meta.get("year", 2026)
    if paras[39].runs:
        paras[39].runs[0].text = f"г. {city}, {year} г."


# ----------------------------------------------------------------------
# Парсер Markdown
# ----------------------------------------------------------------------


_INLINE = re.compile(r"(\*\*[^*]+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def parse_markdown(src: Path) -> tuple[dict[str, Any], list[dict[str, Any]]]:
    text = src.read_text(encoding="utf-8")
    fm: dict[str, Any] = {}
    body = text
    if text.startswith("---\n"):
        end = text.find("\n---", 4)
        if end > 0:
            fm = yaml.safe_load(text[4:end]) or {}
            body = text[end + len("\n---") :]

    blocks: list[dict[str, Any]] = []
    lines = body.splitlines()
    i = 0
    pending_caption: str | None = None
    while i < len(lines):
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("# "):
            blocks.append({"kind": "h1", "text": line[2:].strip()})
            i += 1
            continue
        if line.startswith("## "):
            blocks.append({"kind": "h2", "text": line[3:].strip()})
            i += 1
            continue
        m = re.match(r"^\*\*(Таблица\s+\d+\s*[—–-]\s*.+?)\*\*\s*$", line)
        if m:
            pending_caption = m.group(1).strip()
            i += 1
            continue
        if line.lstrip().startswith("|"):
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 1
            if i < len(lines) and re.match(r"^\s*\|[-: |]+\|\s*$", lines[i]):
                i += 1
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            blocks.append(
                {"kind": "table", "caption": pending_caption, "headers": headers, "rows": rows}
            )
            pending_caption = None
            continue
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not (
            lines[i].startswith("#")
            or lines[i].lstrip().startswith("|")
            or lines[i].lstrip().startswith("**Таблица")
        ):
            para_lines.append(lines[i])
            i += 1
        blocks.append({"kind": "p", "text": " ".join(s.strip() for s in para_lines)})
    return fm, blocks


# ----------------------------------------------------------------------
# Добавление контентного раздела после титульника
# ----------------------------------------------------------------------


def _set_body_run_font(
    run: Any, *, bold: bool | None = None, italic: bool | None = None
) -> None:
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_FONT_SIZE)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), BODY_FONT)


def _add_inline_runs(p: Any, text: str) -> None:
    parts = _INLINE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            _set_body_run_font(run, bold=True)
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            _set_body_run_font(run, italic=True)
        elif part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            _set_body_run_font(run)
            run.font.name = "Consolas"
        else:
            run = p.add_run(part)
            _set_body_run_font(run)


def _set_para_format(p: Any, *, indent: bool, alignment: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pf = p.paragraph_format
    pf.line_spacing = BODY_LINE_SPACING
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if indent:
        pf.first_line_indent = Cm(BODY_INDENT_CM)
    align_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    p.alignment = align_map[alignment]


def _set_cell_borders(cell: Any) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        b = tcBorders.find(qn(f"w:{edge}"))
        if b is None:
            b = OxmlElement(f"w:{edge}")
            tcBorders.append(b)
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "000000")
        b.set(qn("w:space"), "0")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        for row in table.rows:
            for cell in row.cells:
                _set_cell_borders(cell)
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].text = ""
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        _set_body_run_font(run, bold=True)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.paragraphs[0].text = ""
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(cell.paragraphs[0], value)


def append_content(doc: Document, blocks: list[dict[str, Any]]) -> None:
    doc.add_page_break()
    table_counter = 0
    for blk in blocks:
        kind = blk["kind"]
        if kind == "h1":
            p = doc.add_paragraph()
            _set_para_format(p, indent=False, alignment="center")
            run = p.add_run(blk["text"].upper())
            _set_body_run_font(run, bold=True)
        elif kind == "h2":
            p = doc.add_paragraph()
            _set_para_format(p, indent=True, alignment="left")
            run = p.add_run(blk["text"])
            _set_body_run_font(run, bold=True)
        elif kind == "p":
            p = doc.add_paragraph()
            _set_para_format(p, indent=True, alignment="justify")
            _add_inline_runs(p, blk["text"])
        elif kind == "table":
            table_counter += 1
            cap = blk.get("caption") or f"Таблица {table_counter}"
            cap_p = doc.add_paragraph()
            _set_para_format(cap_p, indent=False, alignment="left")
            _add_inline_runs(cap_p, cap)
            _add_table(doc, blk["headers"], blk["rows"])


# ----------------------------------------------------------------------
# main
# ----------------------------------------------------------------------


def build(src: Path, template: Path, out: Path) -> None:
    docx_template = _ensure_docx(template)
    fm, blocks = parse_markdown(src)
    doc = Document(str(docx_template))
    _strip_to_title_page(doc, title_end_paragraph=39)
    fill_title_page_k4(doc, fm)
    append_content(doc, blocks)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("Сохранён %s (%d байт)", out, out.stat().st_size)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--src", type=Path, required=True)
    parser.add_argument("--template", type=Path, required=True)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args(argv)
    if not args.src.exists():
        print(f"Source not found: {args.src}", file=sys.stderr)
        return 2
    if not args.template.exists():
        print(f"Template not found: {args.template}", file=sys.stderr)
        return 2
    build(args.src, args.template, args.out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
